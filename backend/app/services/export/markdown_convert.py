"""Minimal markdown → docx / pdf converters for generated protocols.

Supports the subset of markdown our protocol templates use:
- `#`..`######` headings
- paragraphs separated by blank lines
- `- ` bullet lists
- `| col | col |` tables (one header row + separator + body rows)
- `**bold**` inline emphasis
- `---` horizontal rules

Kept dependency-free beyond python-docx / weasyprint which are already in
requirements.
"""
from __future__ import annotations

import html
import re
from io import BytesIO

_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
_HR_RE = re.compile(r"^-{3,}\s*$")
_BULLET_RE = re.compile(r"^[-*]\s+(.*)$")
_TABLE_ROW_RE = re.compile(r"^\|(.+)\|\s*$")
_TABLE_SEP_RE = re.compile(r"^\|[\s:|-]+\|\s*$")


def _split_table_row(line: str) -> list[str]:
    inner = line.strip().strip("|")
    return [c.strip() for c in inner.split("|")]


# ---------- DOCX ----------

def markdown_to_docx(md: str, title: str | None = None) -> bytes:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Noto Sans"
    style.font.size = Pt(11)

    if title:
        doc.add_heading(title, level=0)

    lines = md.splitlines()
    i = 0
    n = len(lines)

    def add_runs(paragraph, text: str) -> None:
        pos = 0
        for m in _BOLD_RE.finditer(text):
            if m.start() > pos:
                paragraph.add_run(text[pos : m.start()])
            paragraph.add_run(m.group(1)).bold = True
            pos = m.end()
        if pos < len(text):
            paragraph.add_run(text[pos:])

    while i < n:
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if _HR_RE.match(stripped):
            doc.add_paragraph("─" * 40)
            i += 1
            continue

        h = _HEADING_RE.match(stripped)
        if h:
            level = min(len(h.group(1)), 4)
            doc.add_heading(h.group(2).strip(), level=level)
            i += 1
            continue

        # Table: header row + separator row + body rows
        if _TABLE_ROW_RE.match(stripped) and i + 1 < n and _TABLE_SEP_RE.match(lines[i + 1].strip()):
            header = _split_table_row(stripped)
            i += 2
            body_rows: list[list[str]] = []
            while i < n and _TABLE_ROW_RE.match(lines[i].strip()):
                body_rows.append(_split_table_row(lines[i].strip()))
                i += 1
            tbl = doc.add_table(rows=1, cols=len(header))
            try:
                tbl.style = "Light Grid Accent 1"
            except KeyError:
                pass
            hdr = tbl.rows[0].cells
            for idx, cell in enumerate(header):
                hdr[idx].text = cell
            for row in body_rows:
                cells = tbl.add_row().cells
                for idx, val in enumerate(row):
                    if idx < len(cells):
                        cells[idx].text = val
            continue

        b = _BULLET_RE.match(stripped)
        if b:
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, b.group(1))
            i += 1
            continue

        # Plain paragraph (collect until blank line)
        buf = [stripped]
        i += 1
        while i < n and lines[i].strip() and not _HEADING_RE.match(lines[i].strip()) \
                and not _BULLET_RE.match(lines[i].strip()) \
                and not _TABLE_ROW_RE.match(lines[i].strip()) \
                and not _HR_RE.match(lines[i].strip()):
            buf.append(lines[i].strip())
            i += 1
        p = doc.add_paragraph()
        add_runs(p, " ".join(buf))

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------- PDF ----------

def _md_to_html(md: str) -> str:
    out: list[str] = []
    lines = md.splitlines()
    i = 0
    n = len(lines)

    def inline(text: str) -> str:
        escaped = html.escape(text)
        return _BOLD_RE.sub(r"<strong>\1</strong>", escaped)

    while i < n:
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if _HR_RE.match(stripped):
            out.append("<hr/>")
            i += 1
            continue

        h = _HEADING_RE.match(stripped)
        if h:
            level = min(len(h.group(1)), 6)
            out.append(f"<h{level}>{inline(h.group(2).strip())}</h{level}>")
            i += 1
            continue

        if _TABLE_ROW_RE.match(stripped) and i + 1 < n and _TABLE_SEP_RE.match(lines[i + 1].strip()):
            header = _split_table_row(stripped)
            i += 2
            body_rows: list[list[str]] = []
            while i < n and _TABLE_ROW_RE.match(lines[i].strip()):
                body_rows.append(_split_table_row(lines[i].strip()))
                i += 1
            th = "".join(f"<th>{inline(c)}</th>" for c in header)
            trs = "".join(
                "<tr>" + "".join(f"<td>{inline(c)}</td>" for c in r) + "</tr>" for r in body_rows
            )
            out.append(f"<table><thead><tr>{th}</tr></thead><tbody>{trs}</tbody></table>")
            continue

        if _BULLET_RE.match(stripped):
            items: list[str] = []
            while i < n and _BULLET_RE.match(lines[i].strip()):
                items.append(f"<li>{inline(_BULLET_RE.match(lines[i].strip()).group(1))}</li>")
                i += 1
            out.append("<ul>" + "".join(items) + "</ul>")
            continue

        buf = [stripped]
        i += 1
        while i < n and lines[i].strip() and not _HEADING_RE.match(lines[i].strip()) \
                and not _BULLET_RE.match(lines[i].strip()) \
                and not _TABLE_ROW_RE.match(lines[i].strip()) \
                and not _HR_RE.match(lines[i].strip()):
            buf.append(lines[i].strip())
            i += 1
        out.append(f"<p>{inline(' '.join(buf))}</p>")

    return "\n".join(out)


_PDF_CSS = """
@page { size: A4; margin: 20mm; }
body { font-family: "Noto Sans", "DejaVu Sans", sans-serif; font-size: 11pt; color: #111; }
h1 { font-size: 18pt; margin-top: 0; }
h2 { font-size: 14pt; margin-top: 1em; }
h3 { font-size: 12pt; margin-top: 0.8em; }
p  { margin: 0.4em 0; }
ul { margin: 0.2em 0 0.6em 1.2em; }
table { border-collapse: collapse; width: 100%; margin: 0.6em 0; }
th, td { border: 1px solid #888; padding: 4px 6px; text-align: left; vertical-align: top; }
th { background: #f2f2f2; }
hr { border: 0; border-top: 1px solid #bbb; margin: 0.8em 0; }
"""


def markdown_to_pdf(md: str, title: str | None = None) -> bytes:
    from weasyprint import CSS, HTML

    body = _md_to_html(md)
    doc_title = html.escape(title or "Протокол")
    full = (
        f"<!doctype html><html><head><meta charset='utf-8'><title>{doc_title}</title></head>"
        f"<body>{body}</body></html>"
    )
    return HTML(string=full).write_pdf(stylesheets=[CSS(string=_PDF_CSS)])


__all__ = ["markdown_to_docx", "markdown_to_pdf"]
