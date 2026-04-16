"""python-docx exporter."""
from __future__ import annotations

from io import BytesIO
from typing import Any

from app.services.export.formatting import (
    ms_to_clock,
    result_title,
    speaker_label,
    speakers_by_id,
)
from app.services.export.html_template import LABELS, _pick_lang  # reuse localized labels


def render_docx(result: dict[str, Any]) -> bytes:
    from docx import Document
    from docx.shared import Pt

    lang = _pick_lang(result)
    L = LABELS[lang]
    proto = result.get("protocol") or {}
    meta = result.get("metadata") or {}
    spk_map = speakers_by_id(result)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Noto Sans"
    style.font.size = Pt(11)

    title = result_title(result, {"ru": "Протокол заседания", "kk": "Жиналыс хаттамасы", "en": "Meeting minutes"}[lang])
    doc.add_heading(title, level=0)

    date = proto.get("date")
    if date:
        doc.add_paragraph(f"{L['date']}: {date}")
    duration_ms = int(meta.get("duration_ms") or 0)
    doc.add_paragraph(f"{L['duration']}: {ms_to_clock(duration_ms)}")
    doc.add_paragraph(f"{L['languages']}: {', '.join(meta.get('languages_detected') or []) or '—'}")

    # Participants
    doc.add_heading(L["participants"], level=1)
    participants = proto.get("participants") or []
    if participants:
        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = "Light Grid Accent 1"
        hdr = tbl.rows[0].cells
        hdr[0].text = L["p_label"]
        hdr[1].text = L["p_role"]
        for p in participants:
            row = tbl.add_row().cells
            row[0].text = p.get("label") or p.get("id") or ""
            row[1].text = p.get("role") or "—"
    else:
        doc.add_paragraph("—")

    agenda = proto.get("agenda") or []
    if agenda:
        doc.add_heading(L["agenda"], level=1)
        for i, item in enumerate(agenda, 1):
            doc.add_paragraph(f"{i}. {item}")

    discussion = proto.get("discussion") or []
    if discussion:
        doc.add_heading(L["discussion"], level=1)
        for d in discussion:
            p = doc.add_paragraph()
            p.add_run((d.get("topic") or "")).bold = True
            if d.get("summary"):
                doc.add_paragraph(d["summary"])
            if d.get("speakers"):
                doc.add_paragraph(f"{L['speakers']}: {', '.join(d['speakers'])}").italic = True

    decisions = proto.get("decisions") or []
    if decisions:
        doc.add_heading(L["decisions"], level=1)
        for i, dec in enumerate(decisions, 1):
            votes = dec.get("votes") or {}
            tail = ""
            if votes:
                tail = (
                    f" ({L['v_for']}: {votes.get('for', 0)} · "
                    f"{L['v_against']}: {votes.get('against', 0)} · "
                    f"{L['v_abstain']}: {votes.get('abstain', 0)})"
                )
            doc.add_paragraph(f"{i}. {dec.get('text', '')}{tail}")

    actions = proto.get("action_items") or []
    if actions:
        doc.add_heading(L["actions"], level=1)
        tbl = doc.add_table(rows=1, cols=3)
        tbl.style = "Light Grid Accent 1"
        hdr = tbl.rows[0].cells
        hdr[0].text = L["a_task"]
        hdr[1].text = L["a_assignee"]
        hdr[2].text = L["a_deadline"]
        for a in actions:
            row = tbl.add_row().cells
            row[0].text = a.get("task") or ""
            row[1].text = a.get("assignee") or "—"
            row[2].text = a.get("deadline") or "—"

    doc.add_heading(L["transcript"], level=1)
    for s in result.get("transcript") or []:
        label = speaker_label(s.get("speaker", ""), spk_map)
        ts = ms_to_clock(int(s.get("start_time") or 0))
        head = f"[{ts}] {label}"
        if s.get("language"):
            head += f" · {s['language']}"
        p = doc.add_paragraph()
        p.add_run(head).bold = True
        doc.add_paragraph(s.get("text") or "")

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
